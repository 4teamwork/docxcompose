from docx import Document
from utils import ComparableDocument
from utils import docx_path
from utils import FixtureDocument

from docxcompose.sdt import StructuredDocumentTags


def test_set_sdt_text_content():
    doc = Document(docx_path("content_controls.docx"))
    sdt = StructuredDocumentTags(doc)
    sdt.set_text("cc.plain_text", "Foo Bar")
    sdt.set_text("cc.plain_text_multiline", "Foo Bar")
    sdt.set_text("cc.plain_text_empty", "Foo Bar")
    sdt.set_text("cc.rich_text", "Foo Bar")

    updated = ComparableDocument(doc)
    expected = FixtureDocument("content_controls.docx")

    assert updated == expected


def test_set_sdt_multiline_text_content():
    doc = Document(docx_path("content_controls_multiline_formatted.docx"))
    sdt = StructuredDocumentTags(doc)
    sdt.set_text("cc.plain_text", "Line 1\nLine 2")
    sdt.set_text("cc.plain_text_multiline", "Line 1\nLine 2")
    sdt.set_text("cc.plain_text_empty", "Line 1\nLine 2")
    sdt.set_text("cc.rich_text", "Line 1\nLine 2")

    updated = ComparableDocument(doc)
    expected = FixtureDocument("content_controls_multiline_formatted.docx")

    assert updated == expected


def test_get_sdt_multiline_text_content():
    doc = FixtureDocument("content_controls_multiline_formatted.docx")
    sdt = StructuredDocumentTags(doc.doc)

    assert sdt.get_text("cc.plain_text") == "Line 1 Line 2"
    assert sdt.get_text("cc.plain_text_multiline") == "Line 1\nLine 2"
    assert sdt.get_text("cc.plain_text_empty") == "Line 1 Line 2"
