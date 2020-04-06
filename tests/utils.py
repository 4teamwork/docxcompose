from docx import Document
from docxcompose.composer import Composer
from operator import attrgetter
import os.path
from docxcompose.utils import xpath


XPATH_CACHED_DOCPROPERTY_VALUES = 'w:r[preceding-sibling::w:r/w:fldChar/@w:fldCharType="separate"]/w:t'


def docx_path(filename):
    return os.path.join(os.path.dirname(__file__), 'docs', filename)


def simple_field_expression(name):
    return u'.//w:fldSimple[contains(@w:instr, \'DOCPROPERTY "{}"\')]//w:t'.format(name)


def complex_field_expression(name):
    return u'.//w:instrText[contains(.,\'DOCPROPERTY "{}"\')]'.format(name)


def cached_complex_field_values(element):
    value_elements = xpath(element, XPATH_CACHED_DOCPROPERTY_VALUES)
    return [each.text for each in value_elements]


def assert_simple_field_value(expected, element, name):
    prop_elements = xpath(element, simple_field_expression(name))
    assert len(prop_elements) == 1, u'Could not find simple field "{}"'.format(name)
    actual = prop_elements[0].text
    assert expected == actual, u'{} == {}'.format(expected, actual)


def assert_complex_field_value(expected, element, name):
    prop_elements = xpath(element, complex_field_expression(name))
    assert len(prop_elements) == 1, u'Could not find complex field "{}"'.format(name)
    parent_paragraph = prop_elements[0].getparent().getparent()
    actual = u''.join(cached_complex_field_values(parent_paragraph))
    assert expected == actual, u'{} == {}'.format(expected, actual)


class ComparableDocument(object):
    """Test helper to compare two docx documents."""

    def __init__(self, doc):
        self.has_neq_partnames = False
        self.neq_parts = []

        self.doc = doc
        if not doc:
            self.parts = []
            self.partnames = []
            return

        self.parts = sorted(
            self.doc.part.package.parts, key=attrgetter('partname'))
        self.partnames = sorted(p.partname for p in self.parts)

    def __eq__(self, other):
        self.has_neq_partnames = self.partnames != other.partnames
        other.has_neq_partnames = self.has_neq_partnames
        if self.has_neq_partnames:
            return False

        for my_part, other_part in zip(self.parts, other.parts):
            if my_part.blob != other_part.blob:
                self.neq_parts.append((my_part, other_part))
                other.neq_parts.append((other_part, my_part))
        if self.neq_parts:
            return False

        return True

    def post_compare_failed(self, other):
        """Called after a failed comparison/assert."""

        pass


class FixtureDocument(ComparableDocument):
    """Load a comparable document from the composed assets."""

    def __init__(self, composed_filename):
        self.composed_filename = composed_filename

        path = docx_path(os.path.join('composed_fixture', composed_filename))
        doc = Document(path) if os.path.isfile(path) else None

        super(FixtureDocument, self).__init__(doc)


class ComposedDocument(ComparableDocument):
    """Compose at least two documents and provide a docx document for
    comparison.

    Store output document in the `composed_debug` folder when compared to a
    document from the fixture and the assertion failed.

    """
    def __init__(self, master_filename, filename, *filenames):
        composer = Composer(Document(docx_path(master_filename)))
        for filename in (filename,) + filenames:
            composer.append(Document(docx_path(filename)))

        super(ComposedDocument, self).__init__(composer.doc)

    def post_compare_failed(self, other):
        """When comparison to a document from the fixture failed store our
        result in a debug folder.

        """
        if isinstance(other, FixtureDocument):
            path = docx_path(
                os.path.join('composed_debug', other.composed_filename))
            self.doc.save(path)
