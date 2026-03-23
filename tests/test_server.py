from io import BytesIO

import pytest
from aiohttp import FormData
from docx import Document
from utils import ComparableDocument
from utils import docx_path
from utils import FixtureDocument

from docxcompose.server import create_app


@pytest.fixture
async def http_client(aiohttp_client):
    return await aiohttp_client(create_app())


async def test_post_with_documents_returns_composed_document(http_client):
    files = {
        "master": open(docx_path("master.docx"), "rb"),
        "table": open(docx_path("table.docx"), "rb"),
    }
    resp = await http_client.post("/", data=files)
    assert resp.status == 200
    composed_doc = ComparableDocument(Document(BytesIO(await resp.read())))
    composed_fixture = FixtureDocument("table.docx")
    assert composed_doc == composed_fixture


async def test_get_returns_405(http_client):
    resp = await http_client.get("/")
    assert resp.status == 405
    text = await resp.text()
    assert text == "405: Method Not Allowed"


async def test_post_without_multipart_returns_400(http_client):
    resp = await http_client.post("/")
    assert resp.status == 400
    text = await resp.text()
    assert text == "Multipart request required"


async def test_post_returns_400_if_compose_fails(http_client):
    data = FormData()
    data.add_field("master", "foo", content_type="plain/text")
    data.add_field("table", "bar", content_type="plain/text")
    resp = await http_client.post("/", data=data)
    assert resp.status == 400
    text = await resp.text()
    assert text == "No documents provided"


async def test_post_returns_500_if_compose_fails(http_client):
    files = {
        "master": BytesIO(b"FOO"),
        "table": BytesIO(b"bar"),
    }
    resp = await http_client.post("/", data=files)
    assert resp.status == 500
    text = await resp.text()
    assert text == "Failed composing documents"


async def test_post_with_url_parameters(http_client):
    files = {
        "master": open(docx_path("master.docx"), "rb"),
        "table": open(docx_path("table.docx"), "rb"),
    }
    resp = await http_client.post("/?preserve_styles=1", data=files)
    assert resp.status == 200
    composed_doc = ComparableDocument(Document(BytesIO(await resp.read())))
    composed_fixture = FixtureDocument("table.docx")
    assert composed_doc == composed_fixture


async def test_healtcheck_returns_200(http_client):
    resp = await http_client.get("/healthcheck")
    assert resp.status == 200
    text = await resp.text()
    assert text == "OK"
