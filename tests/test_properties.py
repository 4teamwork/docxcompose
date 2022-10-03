from datetime import datetime
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docxcompose.properties import ComplexField
from docxcompose.properties import CUSTOM_PROPERTY_TYPES
from docxcompose.properties import CustomProperties
from docxcompose.properties import SimpleField
from docxcompose.properties import value2vt
from docxcompose.properties import vt2value
from docxcompose.utils import xpath
from lxml.etree import tostring
from utils import assert_complex_field_value
from utils import assert_simple_field_value
from utils import cached_complex_field_values
from utils import docx_path
from utils import simple_field_expression
import pytest


class TestIdentifyDocpropertiesInDocument(object):

    def test_identifies_simple_fields_correctly(self):
        document = Document(docx_path('outdated_docproperty_with_umlauts.docx'))
        properties = CustomProperties(document).find_docprops_in_document()

        assert 1 == len(properties), \
            'input should contain 1 simple field docproperty'

        prop = properties[0]
        assert prop.name == u'F\xfc\xfc'
        assert isinstance(prop, SimpleField)

    def test_identifies_complex_fields_correctly(self):
        document = Document(docx_path('three_props_in_same_paragraph.docx'))
        properties = CustomProperties(document).find_docprops_in_document()

        assert len(document.paragraphs) == 1, 'input file should contains one paragraph'
        assert len(properties) == 3, \
            'input should contain three complex field docproperties'

        # check that all fields were identified as complex fields
        for prop in properties:
            assert isinstance(prop, ComplexField)

        # check that all field names were parsed correctly
        expected_names = ('Text Property', 'Number Property', 'Text Property')
        for name, prop in zip(expected_names, properties):
            assert name == prop.name

        # check that begin, separate and end were identified correctly
        attrib_key = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType'
        expected_indexes = ((1, 3, 11), (13, 15, 17), (25, 27, 32))
        for prop, indexes in zip(properties, expected_indexes):
            assert prop.begin_run.getchildren()[1].attrib[attrib_key] == "begin"
            assert prop.get_separate_run().getchildren()[1].attrib[attrib_key] == "separate"
            assert prop.end_run.getchildren()[1].attrib[attrib_key] == "end"
            assert prop.w_p.index(prop.begin_run) == indexes[0]
            assert prop.w_p.index(prop.get_separate_run()) == indexes[1]
            assert prop.w_p.index(prop.end_run) == indexes[2]

    def test_finds_run_nodes_in_complex_fields_correctly(self):
        document = Document(docx_path('three_props_in_same_paragraph.docx'))
        properties = CustomProperties(document).find_docprops_in_document()

        assert len(properties) == 3, \
            'input should contain three complex field docproperties'

        # In the first field, there are the following runs: begin, docprop,
        # a separate, 3 runs for the value (because of spellcheck) and end
        prop = properties[0]
        assert len(prop.get_runs_for_update()) == 3

        # when dissolving a property the runs that have to be removed
        # are the begin, docprop, and from separate to end (here just
        # separate and end)
        runs = prop.get_runs_to_replace_field_with_value()
        assert len(runs) == 4
        assert runs[0] == prop.begin_run
        assert runs[1] == prop.w_r
        assert runs[2] == prop.get_separate_run()
        assert runs[-1] == prop.end_run

        # In the second field, there are the following runs: begin, docprop,
        # a separate, 1 run for the value and end
        prop = properties[1]
        assert len(prop.get_runs_for_update()) == 1

        # when dissolving a property the runs that have to be removed
        # are from begin to separate (here begin, docprop, separate)
        # and the end
        runs = prop.get_runs_to_replace_field_with_value()
        assert len(runs) == 4
        assert runs[0] == prop.begin_run
        assert runs[1] == prop.w_r
        assert runs[2] == prop.get_separate_run()
        assert runs[-1] == prop.end_run

        # In the first field, there are the following runs: begin, docprop,
        # a separate, 2 runs for the value (because of spellcheck) and end
        prop = properties[2]
        assert len(prop.get_runs_for_update()) == 2

        # when dissolving a property the runs that have to be removed
        # are the begin, docprop, and from separate to end (here just
        # separate and end)
        runs = prop.get_runs_to_replace_field_with_value()
        assert len(runs) == 4
        assert runs[0] == prop.begin_run
        assert runs[1] == prop.w_r
        assert runs[2] == prop.get_separate_run()
        assert runs[-1] == prop.end_run

    def test_finds_run_nodes_in_complex_field_without_separate_correctly(self):
        document = Document(docx_path('complex_field_without_separate.docx'))
        properties = CustomProperties(document).find_docprops_in_document()

        assert len(properties) == 2, \
            'input should contain two complex field docproperties'

        # The "User.FullName" docproperty should be the one without a separate run
        # In this field, there are the following runs: begin, docprop and end
        matches = [prop for prop in properties if prop.name == 'User.FullName']
        assert len(matches) == 1, \
            "There should be only one User.FullName docproperty"
        prop = matches[0]
        assert prop.get_separate_run() is None, \
            "This complex field should not have a separate run."
        assert prop.get_runs_for_update() == [], \
            "As there is no separate run, there should be no run to update"

        # As there are no separate, all runs should be removed when dissolving
        # the property.
        runs = prop.get_runs_to_replace_field_with_value()
        assert len(runs) == 3
        assert runs[0] == prop.begin_run
        assert runs[1] == prop.w_r
        assert runs[2] == prop.end_run

    def test_finds_field_name_and_format_when_split_in_several_runs(self):
        document = Document(docx_path('complex_field_with_split_fieldname.docx'))
        properties = CustomProperties(document).find_docprops_in_document()

        assert len(properties) == 1, \
            'input should contain one complex field docproperties'

        prop = properties[0]

        # make sure that the field name and format are indeed split over several runs
        assert [[each.text for each in xpath(run, "w:instrText")]
                for run in prop._runs[:9]] == \
               [[' DOCPROPERTY "ogg.document.document_date"  \\@ "ddd'],
                ['d'],
                [' dd'],
                [' '],
                ['MM'],
                ['M'],
                ['M'],
                [' yy'],
                ['yy hh:mm:s" \\* MERGEFORMAT ']]

        assert prop._get_fieldname_string() == \
            ' DOCPROPERTY "ogg.document.document_date"  \\@ "ddd' +\
            'd dd MMMM yyyy hh:mm:s" \\* MERGEFORMAT '

        assert prop.name == 'ogg.document.document_date'
        assert prop.date_format == '%A %d %B %Y %H:%M:%-S'

    def test_finds_header_footer_of_different_sections(self):
        document = Document(docx_path('docproperties_header_footer_3_sections.docx'))
        properties = CustomProperties(document).find_docprops_in_document()

        assert len(properties) == 5, \
            'input should contain 5 properties in header/footer and body'

        expected_properties = [
                'Text Property',
                'Number Property',
                'Date Property',
                'Float Property',
                'Boolean Property']

        assert [each.name for each in properties] == expected_properties


class TestUpdateAllDocproperties(object):

    def test_updates_doc_properties_in_header(self):
        document = Document(docx_path('docproperties_header.docx'))

        assert_simple_field_value(
            u'xxx',
            document.sections[0].header.part.element,
            u'my.text-prop')

        CustomProperties(document).update_all()

        assert_simple_field_value(
            u"i'm some text",
            document.sections[0].header.part.element,
            u'my.text-prop')

    def test_updates_doc_properties_in_footer(self):
        document = Document(docx_path('docproperties_footer.docx'))

        assert_simple_field_value(
            u'xxx',
            document.sections[0].footer.part.element,
            u'my.text-prop')

        CustomProperties(document).update_all()

        assert_simple_field_value(
            u'b\xe4hh',
            document.sections[0].footer.part.element,
            u'my.text-prop')

    def test_updates_doc_properties_different_first_page(self):
        document = Document(docx_path('docproperties_different_first_page_1_section.docx'))

        assert_simple_field_value(
            u'xxx',
            document.sections[0].first_page_header.part.element,
            u'page1.header')
        assert_simple_field_value(
            u'xxx',
            document.sections[0].first_page_footer.part.element,
            u'page1.footer')
        assert_simple_field_value(
            u'0',
            document.sections[0].header.part.element,
            u'page2.header')
        assert_simple_field_value(
            u'01.01.1970',
            document.sections[0].footer.part.element,
            u'page2.footer')

        CustomProperties(document).update_all()

        assert_simple_field_value(
            u'p1h',
            document.sections[0].first_page_header.part.element,
            u'page1.header')
        assert_simple_field_value(
            u'p1f',
            document.sections[0].first_page_footer.part.element,
            u'page1.footer')
        assert_simple_field_value(
            u'42',
            document.sections[0].header.part.element,
            u'page2.header')
        assert_simple_field_value(
            u'18.10.1984',
            document.sections[0].footer.part.element,
            u'page2.footer')

    def test_updates_doc_properties_different_odd_even_pages(self):
        document = Document(docx_path('docproperties_different_odd_even_pages_1_section.docx'))

        assert_simple_field_value(
            u'xxx',
            document.sections[0].header.part.element,
            u'odd.header')
        assert_simple_field_value(
            u'xxx',
            document.sections[0].footer.part.element,
            u'odd.footer')
        assert_simple_field_value(
            u'0',
            document.sections[0].even_page_header.part.element,
            u'even.header')
        assert_simple_field_value(
            u'Y',
            document.sections[0].even_page_footer.part.element,
            u'even.footer')

        CustomProperties(document).update_all()

        assert_simple_field_value(
            u'odd-header',
            document.sections[0].header.part.element,
            u'odd.header')
        assert_simple_field_value(
            u'odd-footer',
            document.sections[0].footer.part.element,
            u'odd.footer')
        assert_simple_field_value(
            u'1337',
            document.sections[0].even_page_header.part.element,
            u'even.header')
        assert_simple_field_value(
            u'N',
            document.sections[0].even_page_footer.part.element,
            u'even.footer')

    def test_updates_docproperties_shared_footer(self):
        document = Document(docx_path('docproperties_shared_footer_2_sections.docx'))

        header_1 = document.sections[0].header
        header_2 = document.sections[1].header
        assert not header_1.is_linked_to_previous
        assert not header_2.is_linked_to_previous
        assert_complex_field_value(
            u'xxx',
            header_1.part.element,
            u'section1.header')
        assert_complex_field_value(
            u'yyy',
            header_2.part.element,
            u'section2.header')

        # the same footer should be referenced by both sections
        # the sections should be considered as linked
        footer_1 = document.sections[0].footer
        footer_2 = document.sections[1].footer
        assert not footer_1.is_linked_to_previous
        assert footer_2.is_linked_to_previous
        assert_complex_field_value(
            u'99',
            footer_1.part.element,
            u'shared.footer')
        assert_complex_field_value(
            u'99',
            footer_2.part.element,
            u'shared.footer')

        CustomProperties(document).update_all()

        assert_complex_field_value(
            u'h\xe4der 1',
            header_1.part.element,
            u'section1.header')
        assert_complex_field_value(
            u'h\xf6der 2',
            header_2.part.element,
            u'section2.header')

        # the same footer should be referenced by both sections
        footer_1 = document.sections[0].footer
        footer_2 = document.sections[1].footer
        assert_complex_field_value(
            u'123123123',
            footer_1.part.element,
            u'shared.footer')
        assert_complex_field_value(
            u'123123123',
            footer_2.part.element,
            u'shared.footer')

    def test_updates_docproperties_shared_header(self):
        document = Document(docx_path('docproperties_shared_header_2_sections.docx'))

        # the same header should be referenced by both sections
        # the sections should be considered as linked
        header_1 = document.sections[0].header
        header_2 = document.sections[1].header
        assert not header_1.is_linked_to_previous
        assert header_2.is_linked_to_previous
        assert_complex_field_value(
            u'xxx',
            header_1.part.element,
            u'shared.header')
        assert_complex_field_value(
            u'xxx',
            header_2.part.element,
            u'shared.header')

        footer_1 = document.sections[0].footer
        footer_2 = document.sections[1].footer
        assert not footer_1.is_linked_to_previous
        assert not footer_2.is_linked_to_previous
        assert_complex_field_value(
            u'123',
            footer_1.part.element,
            u'section1.footer')
        assert_complex_field_value(
            u'yyy',
            footer_2.part.element,
            u'section2.footer')

        CustomProperties(document).update_all()

        # the same header should be referenced by both sections
        assert_complex_field_value(
            u'sh\xe4red',
            header_1.part.element,
            u'shared.header')
        assert_complex_field_value(
            u'sh\xe4red',
            header_2.part.element,
            u'shared.header')

        footer_1 = document.sections[0].footer
        footer_2 = document.sections[1].footer
        assert_complex_field_value(
            u'-1.0',
            footer_1.part.element,
            u'section1.footer')
        assert_complex_field_value(
            u'f\xfc\xfcter',
            footer_2.part.element,
            u'section2.footer')

    def test_updates_docproperties_shared_header_footer(self):
        document = Document(docx_path('docproperties_shared_header_footer_2_sections.docx'))

        # the same header should be referenced by both sections
        # the sections should be considered as linked
        header_1 = document.sections[0].header
        header_2 = document.sections[1].header
        assert not header_1.is_linked_to_previous
        assert header_2.is_linked_to_previous
        assert_complex_field_value(
            u'blub',
            header_1.part.element,
            u'shared.header')
        assert_complex_field_value(
            u'blub',
            header_2.part.element,
            u'shared.header')

        # the same footer should be referenced by both sections
        # the sections should be considered as linked
        footer_1 = document.sections[0].footer
        footer_2 = document.sections[1].footer
        assert not footer_1.is_linked_to_previous
        assert footer_2.is_linked_to_previous
        assert_complex_field_value(
            u'N',
            footer_1.part.element,
            u'shared.footer')
        assert_complex_field_value(
            u'N',
            footer_2.part.element,
            u'shared.footer')

        CustomProperties(document).update_all()

        # the same header should be referenced by both sections
        assert_complex_field_value(
            u'ig bi obe',
            header_1.part.element,
            u'shared.header')
        assert_complex_field_value(
            u'ig bi obe',
            header_2.part.element,
            u'shared.header')

        # the same footer should be referenced by both sections
        footer_1 = document.sections[0].footer
        footer_2 = document.sections[1].footer
        assert_complex_field_value(
            u'Y',
            footer_1.part.element,
            u'shared.footer')
        assert_complex_field_value(
            u'Y',
            footer_2.part.element,
            u'shared.footer')

    def test_updates_doc_properties_with_umlauts(self):
        document = Document(docx_path('outdated_docproperty_with_umlauts.docx'))

        assert_simple_field_value(
            u'xxx', document.element.body, u"F\xfc\xfc")

        CustomProperties(document).update_all()

        assert_simple_field_value(
            u'j\xe4ja.', document.element.body, u"F\xfc\xfc")

    def test_complex_docprop_fields_with_multiple_textnodes_are_updated(self):
        document = Document(docx_path('spellchecked_docproperty.docx'))
        paragraphs = xpath(document.element.body, '//w:p')
        assert len(paragraphs) == 1, 'input file contains one paragraph'
        assert len(xpath(document.element.body, '//w:instrText')) == 1, \
            'input contains one complex field docproperty'
        w_p = paragraphs[0]

        cached_values = cached_complex_field_values(w_p)
        assert len(cached_values) == 4, \
            'doc property value is scattered over 4 parts'
        assert ''.join(cached_values) == 'i will be spllchecked!'

        CustomProperties(document).update_all()

        w_p = xpath(document.element.body, '//w:p')[0]
        cached_values = cached_complex_field_values(w_p)
        assert len(cached_values) == 1, \
            'doc property value has been reset to one cached value'
        assert cached_values[0] == 'i will be spllchecked!'

    def test_complex_docprop_with_multiple_textnode_in_same_run_are_updated(self):
        document = Document(docx_path('two_textnodes_in_run_docproperty.docx'))
        paragraphs = xpath(document.element.body, '//w:p')
        assert len(paragraphs) == 1, 'input file contains one paragraph'
        assert len(xpath(document.element.body, '//w:instrText')) == 1, \
            'input contains one complex field docproperty'

        w_p = paragraphs[0]
        cached_values = cached_complex_field_values(w_p)
        assert len(cached_values) == 2, \
            'doc property value is scattered over 2 parts'
        assert ''.join(cached_values) == 'Hello there'

        CustomProperties(document).update_all()

        w_p = xpath(document.element.body, '//w:p')[0]
        cached_values = cached_complex_field_values(w_p)
        assert len(cached_values) == 1, \
            'doc property value has been reset to one cached value'
        assert cached_values[0] == 'i will be spllchecked!'

    def test_three_complex_docprop_in_same_paragraph(self):
        document = Document(docx_path('three_props_in_same_paragraph.docx'))
        properties = CustomProperties(document)

        assert len(document.paragraphs) == 1, 'input file should contains one paragraph'
        paragraph = document.paragraphs[0]
        assert len(properties.find_docprops_in_document()) == 3, \
            'input should contain three complex field docproperties'

        text = u'{text} / {num} mor between the fields {text} and some afte the three fields'
        assert paragraph.text == text.format(text="I was spellcecked", num=0)

        properties.update_all()

        assert paragraph.text == text.format(text="Foo", num=2)

    def test_multiple_identical_docprops_get_updated(self):
        document = Document(docx_path('multiple_identical_properties.docx'))
        assert len(document.paragraphs) == 3, 'input file should contain 3 paragraphs'
        for paragraph in document.paragraphs:
            assert len(xpath(paragraph._p, './/w:instrText')) == 1, \
                'paragraph should contain one complex field docproperties'

            assert paragraph.text == u'Foo'

        CustomProperties(document).update_all()

        for i, paragraph in enumerate(document.paragraphs):
            assert paragraph.text == u'Bar', 'docprop {} was not updated'.format(i+1)

    def test_docproperty_without_separate_does_get_updated(self):
        document = Document(docx_path('complex_field_without_separate.docx'))
        custom_properties = CustomProperties(document)
        properties = custom_properties.find_docprops_in_document()
        paragraphs = document.paragraphs

        # Make sure that a value is set for 'User.FullName'
        assert ('User.FullName', 'Test User') in custom_properties.items()
        assert ('Dossier.Title', ' Some Title') in custom_properties.items()

        # Make sure that 'User.FullName' field has no separate node
        matches = [prop for prop in properties if prop.name == 'User.FullName']
        assert len(matches) == 1, \
            "There should be only one User.FullName docproperty"
        fullname = matches[0]
        assert fullname.get_separate_run() is None, \
            "This complex field should not have a separate run."

        # Make sure that 'Dossier.Title' field has a separate node
        matches = [prop for prop in properties if prop.name == 'Dossier.Title']
        assert len(matches) == 1, \
            "There should be only one Dossier.Title docproperty"
        title = matches[0]
        assert title.get_separate_run() is not None, \
            "This complex field should have a separate run."

        # Check the content of the paragraphs before update
        assert len(paragraphs) == 2
        assert paragraphs[0].text == u'Sachbearbeiter: '
        assert u'Sachbearbeiter: ' in fullname.w_p.xml
        assert paragraphs[1].text == u'Dossier Titel:  '

        custom_properties.update_all()

        # Field with missing separate was not updated
        assert paragraphs[0].text == u'Sachbearbeiter: Test User'
        # Next field was updated correctly
        assert paragraphs[1].text == u'Dossier Titel:  Some Title'

    def test_date_docprops_with_format_get_updated(self):
        document = Document(docx_path('date_docproperties_with_format.docx'))
        assert len(document.paragraphs) == 3, 'input file should contain 3 paragraph'

        expected_values = [u'11.06.19', u'mardi 11 juin 2019', u'11-6-19 0:0:0']
        for i, (expected, paragraph) in enumerate(zip(expected_values, document.paragraphs)):
            assert paragraph.text == expected

        CustomProperties(document).update_all()

        expected_values = [u'23.01.20', u'Thursday 23 January 2020', u'23-1-20 10:0:0']
        for i, (expected, paragraph) in enumerate(zip(expected_values, document.paragraphs)):
            assert paragraph.text == expected, 'docprop {} was not updated correctly'.format(i+1)

    def test_docprops_with_split_fieldname_get_updated(self):
        document = Document(docx_path('complex_field_with_split_fieldname.docx'))
        assert len(document.paragraphs) == 1, 'input file should contain 1 paragraph'

        paragraph = document.paragraphs[0]
        assert paragraph.text == 'Datum: Tuesday 09 February 2021 00:00:00'

        CustomProperties(document).update_all()
        assert paragraph.text == 'Datum: Friday 11 March 2022 10:00:0'


class TestUpdateSpecificDocproperty(object):

    def test_simple_field_gets_updated(self):
        document = Document(docx_path('outdated_docproperty_with_umlauts.docx'))
        assert_simple_field_value(
            u'xxx', document.element.body, u"F\xfc\xfc")

        CustomProperties(document).update(u"F\xfc\xfc", u"new v\xe4lue")

        assert_simple_field_value(
            u"new v\xe4lue", document.element.body, u"F\xfc\xfc")

    def test_complex_field_gets_updated(self):
        document = Document(docx_path('docproperties.docx'))
        assert len(document.paragraphs) == 6, 'input file should contain 6 paragraphs'

        properties = xpath(document.element.body, './/w:instrText')
        assert len(properties) == 5,\
            'input should contain five complex field docproperties'

        expected_paragraphs = [u'Custom Doc Properties',
                               u'Text: Foo Bar',
                               u'Number: 123',
                               u'Boolean: Y',
                               u'Date: 11.06.2019',
                               u'Float: 1.1']
        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

        CustomProperties(document).update("Number Property", 423)

        expected_paragraphs[2] = u'Number: 423'
        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

    def test_multiple_identical_docprops_get_updated(self):
        document = Document(docx_path('multiple_identical_properties.docx'))
        assert len(document.paragraphs) == 3, 'input file should contain 3 paragraphs'
        for paragraph in document.paragraphs:
            assert len(xpath(paragraph._p, './/w:instrText')) == 1, \
                'paragraph should contain one complex field docproperties'

            assert paragraph.text == u'Foo'

        CustomProperties(document).update("Text Property", "New value")

        for i, paragraph in enumerate(document.paragraphs):
            assert paragraph.text == u'New value',\
                'docprop {} was not updated'.format(i+1)


class TestDissolveField(object):

    def test_removes_simple_field_but_keeps_value(self):
        document = Document(docx_path('outdated_docproperty_with_umlauts.docx'))
        assert len(document.paragraphs) == 1, 'input file should contain 1 paragraph'
        fields = xpath(
            document.element.body,
            simple_field_expression(u"F\xfc\xfc"))
        assert len(fields) == 1, 'should contain one simple field docproperty'

        assert document.paragraphs[0].text == u'Hie chund ds property: '
        assert fields[0].text == u'xxx'

        CustomProperties(document).dissolve_fields(u"F\xfc\xfc")
        fields = xpath(
            document.element.body,
            simple_field_expression(u"F\xfc\xfc"))
        assert len(fields) == 0, 'should not contain any docproperties anymore'
        # when simple field is removed, the value is moved one up in the hierarchy
        assert document.paragraphs[0].text == u'Hie chund ds property: xxx'

    def test_removes_complex_field_but_keeps_value(self):
        # test fails because field has 2 spaces before docprop name
        document = Document(docx_path('docproperties.docx'))
        assert len(document.paragraphs) == 6, 'input file should contain 6 paragraphs'

        properties = xpath(document.element.body, './/w:instrText')
        assert len(properties) == 5,\
            'input should contain five complex field docproperties'

        expected_paragraphs = [u'Custom Doc Properties',
                               u'Text: Foo Bar',
                               u'Number: 123',
                               u'Boolean: Y',
                               u'Date: 11.06.2019',
                               u'Float: 1.1']
        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

        CustomProperties(document).dissolve_fields("Number Property")

        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

        properties = xpath(document.element.body, './/w:instrText')
        assert len(properties) == 4,\
            'only 4 fields should remain after removal of one'

    def test_dissolves_all_instances_of_given_field(self):
        document = Document(docx_path('multiple_identical_properties.docx'))
        assert len(document.paragraphs) == 3, 'input file should contain 3 paragraphs'
        assert len(xpath(document.element.body, './/w:instrText')) == 3, \
            'document should contain three complex field docproperties'

        for paragraph in document.paragraphs:
            assert paragraph.text == u'Foo'

        CustomProperties(document).dissolve_fields("Text Property")

        assert len(document.paragraphs) == 3
        assert len(xpath(document.element.body, './/w:instrText')) == 0, \
            'document should not contain any complex field anymore'
        for paragraph in document.paragraphs:
            assert paragraph.text == u'Foo', "value should have been kept in document"

    def test_dissolving_field_when_three_complex_docprop_in_same_paragraph(self):
        document = Document(docx_path('three_props_in_same_paragraph.docx'))
        assert len(document.paragraphs) == 1, 'input file should contains one paragraph'
        paragraph = document.paragraphs[0]
        properties = CustomProperties(document)
        assert len(properties.find_docprops_in_document()) == 3, \
            'input should contain three complex field docproperties'

        text = u'{text} / {num} mor between the fields {text} and some afte the three fields'
        assert paragraph.text == text.format(text="I was spellcecked", num=0)

        properties.dissolve_fields("Text Property")

        assert len(document.paragraphs) == 1
        assert len(properties.find_docprops_in_document()) == 1, \
            'document should contain one complex field after removal'
        assert paragraph.text == text.format(text="I was spellcecked", num=0)


def test_get_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props['Text Property'] == 'Foo Bar'
    assert props['Number Property'] == 123
    assert props['Boolean Property'] is True
    assert props['Date Property'] == datetime(2019, 6, 11, 10, 0)

    assert props.get('Text Property') == 'Foo Bar'
    assert props.get('Number Property') == 123
    assert props.get('Boolean Property') is True
    assert props.get('Date Property') == datetime(2019, 6, 11, 10, 0)


def test_get_doc_property_is_case_insensitive():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props['text property'] == 'Foo Bar'
    assert props.get('text property') == 'Foo Bar'


def test_add_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props.add('My Text Property', 'foo bar')
    assert props.get('My Text Property') == 'foo bar'

    props.add('My Boolean Property', True)
    assert props.get('My Boolean Property') is True

    props.add('My Number Property', 123)
    assert props.get('My Number Property') == 123

    props.add('My Date Property', datetime(2019, 10, 23, 15, 44, 50))
    assert props.get('My Date Property') == datetime(2019, 10, 23, 15, 44, 50)


def test_add_utf8_property():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props.add('My Text Property', u'f\xfc\xfc'.encode('utf-8'))
    assert props.get('My Text Property') == u'f\xfc\xfc'


def test_set_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props['Text Property'] = 'baz'
    assert props['Text Property'] == 'baz'

    props['Boolean Property'] = False
    assert props['Boolean Property'] is False

    props['Number Property'] = 456
    assert props['Number Property'] == 456

    props['Date Property'] = datetime(2019, 10, 20, 12, 0)
    assert props['Date Property'] == datetime(2019, 10, 20, 12, 0)


def test_set_doc_property_is_case_insensitive():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props['text property'] = 'baz'
    assert props['Text Property'] == 'baz'


def test_delete_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    del props['Text Property']
    del props['Number Property']

    assert 'Text Property' not in props
    assert 'Number Property' not in props

    assert xpath(props._element, u'.//cp:property/@pid') == ['2', '3', '4']


def test_delete_doc_property_is_case_insensitive():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    del props['text property']

    assert 'Text Property' not in props


def test_contains_is_case_insensitive():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert 'text property' in props


def test_nullify_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props.nullify('Text Property')
    props.nullify('Number Property')
    props.nullify('Boolean Property')
    props.nullify('Date Property')
    props.nullify('Float Property')

    assert props['Text Property'] == ''
    assert 'Number Property' not in props
    assert 'Boolean Property' not in props
    assert 'Date Property' not in props
    assert 'Float Property' not in props


def test_nullify_doc_property_is_case_insensitive():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props.nullify('text property')
    assert props['Text Property'] == ''


def test_set_doc_property_on_document_without_properties_creates_new_part():
    document = Document(docx_path('master.docx'))
    props = CustomProperties(document)
    props['Text Property'] = 'Foo'

    assert props.part is not None
    assert props['Text Property'] == 'Foo'

    part = document.part.package.part_related_by(RT.CUSTOM_PROPERTIES)
    assert part is not None


def test_doc_properties_keys():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props.keys() == [
        'Text Property',
        'Number Property',
        'Boolean Property',
        'Date Property',
        'Float Property',
    ]


def test_doc_properties_values():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props.values() == [
        'Foo Bar', 123, True, datetime(2019, 6, 11, 10, 0), 1.1]


def test_doc_properties_items():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props.items() == [
        ('Text Property', 'Foo Bar'),
        ('Number Property', 123),
        ('Boolean Property', True),
        ('Date Property', datetime(2019, 6, 11, 10, 0)),
        ('Float Property', 1.1),
    ]


def test_vt2value_value2vt_roundtrip():
    assert vt2value(value2vt(42)) == 42
    assert vt2value(value2vt(True)) is True
    assert vt2value(value2vt(1.1)) == pytest.approx(1.1)
    dt = datetime(2019, 6, 11, 10, 0)
    assert vt2value(value2vt(dt)) == dt
    assert vt2value(value2vt(u'foo')) == u'foo'
    assert vt2value(value2vt(u'')) == u''

    node = parse_xml(CUSTOM_PROPERTY_TYPES['int'])
    node.text = '42'
    assert tostring(value2vt(vt2value(node))) == tostring(node)

    node = parse_xml(CUSTOM_PROPERTY_TYPES['bool'])
    node.text = 'true'
    assert tostring(value2vt(vt2value(node))) == tostring(node)

    node = parse_xml(CUSTOM_PROPERTY_TYPES['float'])
    node.text = '1.1'
    assert tostring(value2vt(vt2value(node))) == tostring(node)

    node = parse_xml(CUSTOM_PROPERTY_TYPES['datetime'])
    node.text = '2003-12-31T10:14:55Z'
    assert tostring(value2vt(vt2value(node))) == tostring(node)

    node = parse_xml(CUSTOM_PROPERTY_TYPES['text'])
    node.text = 'foo'
    assert tostring(value2vt(vt2value(node))) == tostring(node)

    node = parse_xml(CUSTOM_PROPERTY_TYPES['text'])
    node.text = ''
    assert tostring(value2vt(vt2value(node))) == tostring(node)


def test_vt2value_returns_empty_string_for_missing_text_node():
    node = parse_xml(CUSTOM_PROPERTY_TYPES['text'])
    node.text = None
    assert vt2value(node) == u''
