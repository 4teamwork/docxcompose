import pytest
from docx import Document
from utils import ComparableDocument
from utils import ComposedDocument
from utils import docx_path
from utils import FixtureDocument

from docxcompose.composer import Composer


def test_contains_predefined_styles_in_masters_language(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert "Heading1" in style_ids
    assert "Heading1" in style_ids
    assert "Strong" in style_ids
    assert "Quote" in style_ids


def test_does_not_contain_predefined_styles_in_appended_language(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert "berschrift1" not in style_ids
    assert "berschrift2" not in style_ids
    assert "Fett" not in style_ids
    assert "Zitat" not in style_ids


def test_contains_custom_styles_from_both_docs(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert "MyStyle1" in style_ids
    assert "MyStyle1Char" in style_ids
    assert "MeineFormatvorlage" in style_ids
    assert "MeineFormatvorlageZchn" in style_ids


def test_contains_linked_styles(merged_styles):
    style_ids = [s.style_id for s in merged_styles.doc.styles]
    assert "QuoteChar" in style_ids


def test_merged_styles_de():
    doc = FixtureDocument("styles_de.docx")
    composed = ComposedDocument("styles_de.docx", "styles_en.docx")

    assert composed == doc


def test_merged_styles_en():
    doc = FixtureDocument("styles_en.docx")
    composed = ComposedDocument("styles_en.docx", "styles_de.docx")

    assert composed == doc


def test_styles_are_not_switched_for_first_numbering_element():
    doc = FixtureDocument("switched_listing_style.docx")
    composed = ComposedDocument(
        "master_switched_listing_style.docx", "switched_listing_style.docx"
    )

    assert composed == doc


def test_continue_when_no_styles():
    """Expects not to throw a type error"""
    ComposedDocument("aatmay.docx", "aatmay.docx")


def test_preserve_styles_with_same_id():
    composer = Composer(
        Document(docx_path("styles_preserve1.docx")), preserve_styles=True
    )
    composer.append(Document(docx_path("styles_preserve2.docx")))
    style_ids = [s.style_id for s in composer.doc.styles]
    assert "MyCustomStyle" in style_ids
    assert "MyCustomStyle_1" in style_ids

    expected = FixtureDocument("styles_preserve.docx")
    composed = ComparableDocument(composer.doc)
    assert composed == expected


def test_ignore_styles_with_same_id():
    composer = Composer(Document(docx_path("styles_preserve1.docx")))
    composer.append(Document(docx_path("styles_preserve2.docx")))
    style_ids = [s.style_id for s in composer.doc.styles]
    assert "MyCustomStyle" in style_ids
    assert "MyCustomStyle_1" not in style_ids


def test_preserve_styles_does_not_duplicate_identical_styles():
    composer = Composer(
        Document(docx_path("styles_preserve1.docx")), preserve_styles=True
    )
    composer.append(Document(docx_path("styles_preserve2.docx")))
    composer.append(Document(docx_path("styles_preserve2.docx")))
    composer.append(Document(docx_path("styles_preserve1.docx")))
    assert [
        s.style_id
        for s in composer.doc.styles
        if s.style_id.startswith("MyCustomStyle")
    ] == ["MyCustomStyle", "MyCustomStyleZchn", "MyCustomStyle_1"]


@pytest.fixture
def merged_styles():
    composer = Composer(Document(docx_path("styles_en.docx")))
    composer.append(Document(docx_path("styles_de.docx")))
    return composer
