from docx import Document
from docxcompose.composer import Composer
from docxcompose.utils import xpath
from utils import ComposedDocument
from utils import docx_path
from utils import FixtureDocument
import pytest


def test_header_and_footer_refs_in_paragraph_props_get_removed(header_footer):
    refs = xpath(
        header_footer.doc.element.body,
        './/w:pPr/w:sectPr/w:headerReference|.//w:pPr/w:sectPr/w:footerReference')
    assert len(refs) == 0


def test_master_header_and_footer_are_preserved_when_adding_sections():
    doc = FixtureDocument("master_header_footer_with_sections.docx")
    composed = ComposedDocument(
        "master_header_footer.docx", "header_footer_sections.docx")

    assert composed == doc


def test_header_footer():
    doc = FixtureDocument("header_footer.docx")
    composed = ComposedDocument(
        "header_footer.docx", "header_footer.docx")
    assert composed == doc


def test_header_footer_sections():
    doc = FixtureDocument("header_footer_sections.docx")
    composed = ComposedDocument(
        "header_footer_sections.docx", "header_footer_sections.docx")

    assert composed == doc


@pytest.fixture
def header_footer():
    composer = Composer(Document(docx_path("master.docx")))
    composer.append(Document(docx_path("header_footer_sections.docx")))
    return composer
