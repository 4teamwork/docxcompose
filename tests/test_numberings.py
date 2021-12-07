from docx import Document
from docxcompose.composer import Composer
from docxcompose.utils import xpath
from utils import ComposedDocument
from utils import docx_path
from utils import FixtureDocument
import pytest


def test_abstractnums_from_styles_are_not_duplicated(multiple_numberings):
    anums = xpath(
        multiple_numberings.doc.part.numbering_part.element,
        './/w:abstractNum[.//w:pStyle]')
    assert len(anums) == 2


def test_restart_first_numbering(multiple_numberings):
    paragraphs = xpath(multiple_numberings.doc.element.body, './/w:p')
    assert len(xpath(paragraphs[9], './/w:numId')) == 1


def test_do_not_restart_numbering_of_bullets(mixed_numberings):
    paragraphs = xpath(mixed_numberings.doc.element.body, './/w:p')
    assert len(xpath(paragraphs[10], './/w:numId')) == 0


def test_do_not_break_on_custom_styled_numbering(custom_styled_numbering):
    assert custom_styled_numbering.doc.element.xpath('.//w:numId/@w:val') == ['2']*4


def test_preserve_zero_numbering_references(numberings_with_zero_reference):
    numberings_with_zero_ref = xpath(
        numberings_with_zero_reference.doc.element.body,
        './/w:p//w:numId[@w:val="0"]')
    assert len(numberings_with_zero_ref) == 2


def test_restarts_numbering_if_sequence_is_split_across_elements(numbering_with_paragraphs):
    numbering_ids = [each.val for each in xpath(
        numbering_with_paragraphs.doc.element.body, './/w:numId')]

    assert numbering_ids == [3, 3, 3, 3, 4, 4, 4]


def test_restart_numbering_manages_shared_style_names(static_reseed):
    doc = FixtureDocument("common_stylename_different_id.docx")
    composed = ComposedDocument(
        "common_stylename_different_id1.docx",
        "common_stylename_different_id2.docx")
    assert composed == doc


def test_numberings(static_reseed):
    doc = FixtureDocument("numberings.docx")
    composed = ComposedDocument(
        "numberings.docx", "numberings.docx")

    assert composed == doc


def test_restart_numberings():
    doc = FixtureDocument("numberings_restart.docx")
    composed = ComposedDocument(
        "numberings_restart.docx", "numberings_restart.docx")

    assert composed == doc


def test_numberings_styles():
    doc = FixtureDocument("numberings_styles.docx")
    composed = ComposedDocument(
        "numberings_styles.docx", "numberings_styles.docx")

    assert composed == doc


def test_numbering_reference_to_numbering_zero():
    doc = FixtureDocument("numbering_reference_to_numbering_zero.docx")
    composed = ComposedDocument("numbering_reference_to_numbering_zero.docx",
                                "numbering_reference_to_numbering_zero.docx")

    assert composed == doc


def test_restarts_numbering_for_all_elements_of_same_sequence():
    doc = FixtureDocument("numbering_with_paragraphs_in_between.docx")
    composed = ComposedDocument(
        "numbering_with_paragraphs_in_between.docx",
        "numbering_with_paragraphs_in_between.docx")

    assert composed == doc


def test_preserves_list_styles_when_restarting_numberings():
    doc = FixtureDocument("broken_listing.docx")
    composed = ComposedDocument(
        "broken_listing_master.docx", "broken_listing.docx")

    assert composed == doc


def test_preserves_list_styles_when_restarting_many_numberings():
    doc = FixtureDocument("broken_listing_many.docx")
    composed = ComposedDocument(
        "broken_listing_master.docx", "broken_listing_many.docx")

    assert composed == doc


def test_preserves_list_styles_when_restarting_nested_numberings():
    doc = FixtureDocument("broken_listing_nested.docx")
    composed = ComposedDocument(
        "broken_listing_master.docx", "broken_listing_nested.docx")

    assert composed == doc


@pytest.fixture
def numberings_with_zero_reference():
    composer = Composer(Document(
        docx_path("numbering_reference_to_numbering_zero.docx")))
    composer.append(Document(
        docx_path("numbering_reference_to_numbering_zero.docx")))
    return composer


@pytest.fixture
def numberings_in_styles():
    composer = Composer(Document(docx_path("master.docx")))
    composer.append(Document(docx_path("numberings_styles.docx")))
    return composer


@pytest.fixture
def multiple_numberings():
    composer = Composer(Document(docx_path("numberings_styles.docx")))
    composer.append(Document(docx_path("numberings_styles.docx")))
    return composer


@pytest.fixture
def mixed_numberings():
    composer = Composer(Document(docx_path("numberings_restart.docx")))
    composer.append(Document(docx_path("numberings_restart.docx")))
    return composer


@pytest.fixture
def numbering_with_paragraphs():
    composer = Composer(Document(docx_path("master.docx")))
    composer.append(Document(docx_path("numbering_with_paragraphs_in_between.docx")))
    return composer


@pytest.fixture
def custom_styled_numbering():
    composer = Composer(Document(docx_path('master.docx')))
    composer.append(Document(docx_path("custom_list_style.docx")))
    return composer
